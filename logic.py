import re
import os
import requests
import logging
from docx import Document
from contextlib import contextmanager
from config import SHORT_WORDS
from Date_Spellcheck_Logic import process_paragraph_spellcheck


def yandex_spellcheck(text: str):
    """
    Проверка орфографии с помощью Яндекс.Спеллера
    """
    try:
        response = requests.post(
            'https://yandex.speller.yandex.net/v1/checkText',
            data={
                'text': text,
                'format': 'plain',
                'lang': 'ru',
                'options': 511  # Максимальный уровень проверки
            }
        )
        return response.json()
    except Exception as e:
        logging.error(f"Ошибка при проверке орфографии: {e}")
        return []


@contextmanager
def safe_document_handling(input_path, output_path):
    """Контекстный менеджер для безопасной работы с документом."""
    # Проверка существования файла
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Файл {input_path} не найден.")

    # Проверка, открыт ли файл
    try:
        with open(input_path, 'rb') as _:
            pass
    except PermissionError:
        raise PermissionError(f"Файл {input_path} уже открыт в другой программе.")

    try:
        doc = Document(input_path)
        yield doc
        # Сохраняем файл только если все операции прошли успешно
        doc.save(output_path)
    except Exception as e:
        # Перебрасываем исключение наверх для обработки в UI
        raise e


def fix_dates_in_paragraph(paragraph):
    """
    Заменяет пробелы в датах на неразрывные пробелы.
    Поддерживает форматы: 26.01.1990 и 26 января 1990
    """
    # Регулярные выражения для различных форматов даты
    date_patterns = [
        r'\d{1,2}\.\d{1,2}\.\d{4}',  # дд.мм.гггг
        r'\d{1,2}\s+[а-яА-Я]+\s+\d{4}'  # дд месяц гггг
    ]

    for run in paragraph.runs:
        text = run.text
        for pattern in date_patterns:
            # Находим все даты в тексте
            dates = re.findall(pattern, text)
            for date in dates:
                # Заменяем пробелы на неразрывные пробелы
                text = text.replace(date, date.replace(' ', '\u00A0'))
        run.text = text


def find_hanging_prepositions(paragraph):
    """
    Улучшенная функция для нахождения висячих предлогов, с особым вниманием к коротким предлогам
    """
    replacements = []

    # Сначала анализируем каждый run отдельно
    for i, run in enumerate(paragraph.runs):
        text = run.text

        # Используем более строгий паттерн для предлогов
        # Добавляем проверку границ слов \b с обеих сторон для коротких предлогов
        pattern = r'(?<!\w)(' + '|'.join(re.escape(word) for word in SHORT_WORDS) + r')(?=\s)'

        for match in re.finditer(pattern, text, re.IGNORECASE):
            start, end = match.span()
            # Проверяем, что после предлога идет пробел
            if end < len(text) and text[end].isspace():
                replacements.append((i, start, end + 1))  # +1 чтобы захватить пробел

    # Ищем предлоги на границах runs с улучшенной логикой
    for i in range(len(paragraph.runs) - 1):
        run_text = paragraph.runs[i].text
        next_run_text = paragraph.runs[i + 1].text

        if not run_text or not next_run_text:
            continue

        # Проверяем, заканчивается ли run на предлог
        last_word_match = re.search(r'(?<!\w)(' + '|'.join(re.escape(word) for word in SHORT_WORDS) + r')$', run_text,
                                    re.IGNORECASE)

        if last_word_match:
            # Если следующий run начинается с пробела или невидимого символа форматирования
            if next_run_text and (next_run_text[0].isspace() or ord(next_run_text[0]) < 32):
                replacements.append((i + 1, 0, 1, 'boundary'))

    return replacements


def process_paragraph(paragraph):
    """
    Обрабатывает параграф, сохраняя форматирование
    """
    if not paragraph.runs:
        return

    # Сначала фиксируем даты
    fix_dates_in_paragraph(paragraph)

    # Затем работаем с предлогами
    replacements = find_hanging_prepositions(paragraph)

    # Применяем замены с конца, чтобы не сбивать индексы
    replacements.sort(reverse=True, key=lambda x: (x[0], x[1]))

    for rep in replacements:
        if len(rep) == 4 and rep[3] == 'boundary':  # Замена на границе runs
            i, start, end = rep[0], rep[1], rep[2]
            if i < len(paragraph.runs) and paragraph.runs[i].text:
                # Заменяем пробел на неразрывный в начале следующего run
                run_text = paragraph.runs[i].text
                if run_text[0].isspace():
                    paragraph.runs[i].text = "\u00A0" + run_text[1:]
                else:  # Обработка невидимых символов
                    paragraph.runs[i].text = "\u00A0" + run_text
        else:
            i, start, end = rep[0], rep[1], rep[2]
            if i < len(paragraph.runs):
                run_text = paragraph.runs[i].text
                # Заменяем пробел на неразрывный внутри run
                paragraph.runs[i].text = run_text[:start] + run_text[start:end - 1] + "\u00A0" + run_text[end:]

    # Проверка орфографии должна быть последним шагом
    process_paragraph_spellcheck(paragraph)


def fix_hanging_prepositions(input_path, output_path, progress_callback=None):
    """
    Основная функция обработки документа.

    Args:
        input_path: Путь к исходному файлу
        output_path: Путь для сохранения обработанного файла
        progress_callback: Функция обратного вызова для обновления прогресса
    """
    with safe_document_handling(input_path, output_path) as doc:
        total_elements = (
                len(doc.paragraphs) +
                sum(len(row.cells) for table in doc.tables for row in table.rows) +
                sum(len(section.header.paragraphs) + len(section.footer.paragraphs) for section in doc.sections)
        )

        processed = 0

        # Обрабатываем обычные абзацы
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
            processed += 1
            if progress_callback:
                progress_callback(processed / total_elements)

        # Обрабатываем текст в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
                        processed += 1
                        if progress_callback:
                            progress_callback(processed / total_elements)

        # Обрабатываем колонтитулы
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                process_paragraph(paragraph)
                processed += 1
                if progress_callback:
                    progress_callback(processed / total_elements)
            for paragraph in section.footer.paragraphs:
                process_paragraph(paragraph)
                processed += 1
                if progress_callback:
                    progress_callback(processed / total_elements)

def fix_hanging_prepositions_with_spellcheck(input_path, output_path, progress_callback=None):
    """
    Обработка документа с заменой предлогов и проверкой орфографии.

    Args:
        input_path: Путь к исходному файлу
        output_path: Путь для сохранения обработанного файла
        progress_callback: Функция обратного вызова для обновления прогресса
    """
    with safe_document_handling(input_path, output_path) as doc:
        total_elements = (
                len(doc.paragraphs) +
                sum(len(row.cells) for table in doc.tables for row in table.rows) +
                sum(len(section.header.paragraphs) + len(section.footer.paragraphs) for section in doc.sections)
        )

        processed = 0

        # Обрабатываем обычные абзацы
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
            # Дополнительная проверка орфографии
            process_paragraph_spellcheck(paragraph)
            processed += 1
            if progress_callback:
                progress_callback(processed / total_elements)

        # Обрабатываем текст в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
                        # Дополнительная проверка орфографии
                        process_paragraph_spellcheck(paragraph)
                        processed += 1
                        if progress_callback:
                            progress_callback(processed / total_elements)

        # Обрабатываем колонтитулы
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                process_paragraph(paragraph)
                # Дополнительная проверка орфографии
                process_paragraph_spellcheck(paragraph)
                processed += 1
                if progress_callback:
                    progress_callback(processed / total_elements)

            for paragraph in section.footer.paragraphs:
                process_paragraph(paragraph)
                # Дополнительная проверка орфографии
                process_paragraph_spellcheck(paragraph)
                processed += 1
                if progress_callback:
                    progress_callback(processed / total_elements)
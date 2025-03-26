import re
import requests
import logging
from docx import Document
from typing import List, Dict


def fix_dates_in_paragraph(paragraph):
    """
    Заменяет пробелы в датах на неразрывные пробелы.
    Поддерживает форматы: 26.01.1990 и 26 января 1990
    """
    # Регулярные выражения для различных форматов даты
    date_patterns = [
        r'\d{1,2}\.\d{1,2}\.\d{4}',  # дд.мм.гггг
        r'\d{1,2}\s+[а-яА-Я]+\s+\d{4}'  # дд месяц гггг
    ]

    for run in paragraph.runs:
        text = run.text
        for pattern in date_patterns:
            # Находим все даты в тексте
            dates = re.findall(pattern, text)
            for date in dates:
                # Заменяем пробелы на неразрывные пробелы
                text = text.replace(date, date.replace(' ', '\u00A0'))
        run.text = text


def yandex_spellcheck(text: str):
    """
    Проверка орфографии с помощью Яндекс.Спеллера
    """
    try:
        # Используем GET-запрос вместо POST
        response = requests.get(
            'https://speller.yandex.net/services/spellservice.json/checkText',
            params={
                'text': text,
                'format': 'plain',
                'lang': 'ru',
                'options': 511  # Максимальный уровень проверки
            }
        )
        return response.json()
    except Exception as e:
        logging.error(f"Ошибка при проверке орфографии: {e}")
        return []


def apply_spellcheck_to_run(run, corrections):
    """
    Применяет исправления орфографии к конкретному run
    с сохранением форматирования
    """
    text = run.text
    for error in corrections:
        start = error.get('pos', 0)
        end = start + error.get('len', 0)
        suggestions = error.get('s', [])

        if suggestions:
            # Заменяем слово первым предложенным вариантом
            text = text[:start] + suggestions[0] + text[end:]

    run.text = text


def process_paragraph_spellcheck(paragraph):
    """
    Проверяет орфографию в параграфе
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])

    # Получаем исправления от Яндекс.Спеллера
    corrections = yandex_spellcheck(full_text)

    if corrections:
        # Применяем исправления к каждому run
        for run in paragraph.runs:
            apply_spellcheck_to_run(run, corrections)


def process_document_with_dates_and_spellcheck(input_path, output_path, progress_callback=None):
    """
    Обрабатывает документ: исправляет даты и проверяет орфографию
    """
    doc = Document(input_path)
    total_elements = (
            len(doc.paragraphs) +
            sum(len(row.cells) for table in doc.tables for row in table.rows) +
            sum(len(section.header.paragraphs) + len(section.footer.paragraphs) for section in doc.sections)
    )

    processed = 0

    # Обрабатываем обычные абзацы
    for paragraph in doc.paragraphs:
        fix_dates_in_paragraph(paragraph)
        process_paragraph_spellcheck(paragraph)
        processed += 1
        if progress_callback:
            progress_callback(processed / total_elements)

    # Обрабатываем текст в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_dates_in_paragraph(paragraph)
                    process_paragraph_spellcheck(paragraph)
                    processed += 1
                    if progress_callback:
                        progress_callback(processed / total_elements)

    # Обрабатываем колонтитулы
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            fix_dates_in_paragraph(paragraph)
            process_paragraph_spellcheck(paragraph)
            processed += 1
            if progress_callback:
                progress_callback(processed / total_elements)

        for paragraph in section.footer.paragraphs:
            fix_dates_in_paragraph(paragraph)
            process_paragraph_spellcheck(paragraph)
            processed += 1
            if progress_callback:
                progress_callback(processed / total_elements)

    doc.save(output_path)